from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")


def flush_paragraph(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    text = " ".join(line.strip() for line in lines if line.strip())
    if text:
        document.add_paragraph(text)
    lines.clear()


def add_code_block(document: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    p = document.add_paragraph()
    for idx, line in enumerate(code_lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if idx < len(code_lines) - 1:
            run.add_break(WD_BREAK.LINE)


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        for row in rows:
            document.add_paragraph(" | ".join(row))
        return
    header = rows[0]
    body = [row for row in rows[2:]] if len(rows) >= 3 else []
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for idx, value in enumerate(header):
        table.rows[0].cells[idx].text = value
    for row in body:
        tr = table.add_row().cells
        for idx, value in enumerate(row):
            if idx < len(tr):
                tr[idx].text = value


def add_image(document: Document, md_path: Path, image_path_text: str) -> None:
    image_path = (md_path.parent / image_path_text).resolve()
    if not image_path.exists():
        document.add_paragraph(f"[Imagen no encontrada: {image_path_text}]")
        return
    document.add_picture(str(image_path), width=Inches(6.3))


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

    for line in lines:
        if line.strip().startswith("```"):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            if in_code:
                add_code_block(document, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not line.strip():
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            flush_paragraph(document, paragraph_buffer)
            table_buffer.append(line)
            continue
        else:
            flush_table()

        heading_match = HEADING_RE.match(line)
        if heading_match:
            flush_paragraph(document, paragraph_buffer)
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(heading_match.group(2).strip(), level=level)
            continue

        image_match = IMAGE_RE.search(line.strip())
        if image_match and line.strip().startswith("!["):
            flush_paragraph(document, paragraph_buffer)
            add_image(document, md_path, image_match.group(2))
            continue

        if line.strip().startswith("- "):
            flush_paragraph(document, paragraph_buffer)
            document.add_paragraph(line.strip()[2:].strip(), style="List Bullet")
            continue

        ordered_match = ORDERED_RE.match(line.strip())
        if ordered_match:
            flush_paragraph(document, paragraph_buffer)
            document.add_paragraph(ordered_match.group(1).strip(), style="List Number")
            continue

        paragraph_buffer.append(line)

    flush_paragraph(document, paragraph_buffer)
    flush_table()
    if code_buffer:
        add_code_block(document, code_buffer)

    document.save(str(docx_path))


def main() -> None:
    if len(sys.argv) not in {2, 3}:
        raise SystemExit("Uso: python scripts/markdown_to_docx.py <input.md> [output.docx]")
    md_path = Path(sys.argv[1]).resolve()
    if len(sys.argv) == 3:
        docx_path = Path(sys.argv[2]).resolve()
    else:
        docx_path = md_path.with_suffix(".docx")
    convert_markdown_to_docx(md_path, docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
